import os
import pandas as pd
import tkinter as tk
from tkinter import filedialog, messagebox
import re
from docx import Document

# Supported image extensions
IMAGE_EXTENSIONS = ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.gif']

# Function to clean and normalize strings for comparison
def clean_string(s):
    # Remove non-alphanumeric characters, convert to lowercase
    return re.sub(r'\W+', '', s.lower())

# Function to get all image files in a directory recursively
def get_all_images(root_dir):
    all_images = {}
    for dirpath, _, filenames in os.walk(root_dir):
        for file in filenames:
            if os.path.splitext(file)[1].lower() in IMAGE_EXTENSIONS:
                all_images[file.lower()] = os.path.join(dirpath, file)
    return all_images

# Function to save missing images in a Word file
def save_missing_images_to_word(missing_images, output_directory):
    # Create a new Word document
    doc = Document()
    doc.add_heading('Missing Images Report', level=1)
    
    # Add each missing image name to the document
    for img in missing_images:
        doc.add_paragraph(img)
    
    # Save the document in the output directory
    output_path = os.path.join(output_directory, 'Missing_Images_Report.docx')
    doc.save(output_path)
    return output_path

# Function to save renamed images log to an Excel file
def save_renamed_images_log(renamed_images, output_directory):
    # Create a DataFrame for renamed images
    df = pd.DataFrame(renamed_images, columns=['Original Name', 'New Name', 'File Path'])
    
    # Save the DataFrame to an Excel file
    output_path = os.path.join(output_directory, 'Renamed_Images_Log.xlsx')
    df.to_excel(output_path, index=False)
    return output_path

# Function to rename images based on the concatenated format: prd_code Prd_name number
def rename_images():
    # Get document file path
    doc_file = filedialog.askopenfilename(filetypes=[("Excel Files", "*.xlsx *.xls")])
    if not doc_file:
        return

    # Read the Excel file
    try:
        df = pd.read_excel(doc_file)
    except Exception as e:
        messagebox.showerror("Error", f"Failed to read the document: {str(e)}")
        return

    # Check if required columns are present
    required_columns = ['prd_code', 'Prd_name', 'number']
    for col in required_columns:
        if col not in df.columns:
            messagebox.showerror("Error", f"'{col}' column not found in the document.")
            return

    # Get image directory
    image_directory = filedialog.askdirectory(title="Select Image Directory")
    if not image_directory:
        return

    # Get output directory
    output_directory = filedialog.askdirectory(title="Select Output Directory")
    if not output_directory:
        return

    # Normalize paths
    image_directory = os.path.normpath(image_directory)
    output_directory = os.path.normpath(output_directory)

    # Get all images in the main directory and its subdirectories
    available_images = get_all_images(image_directory)

    # Track renamed and missing images
    renamed_images = []
    missing_images = []

    # Loop through the document list and try to rename matching images
    for index, row in df.iterrows():
        # Extract and clean components
        prd_code = str(row['prd_code']).strip()
        prd_name = str(row['Prd_name']).strip()
        number = str(row['number']).strip()

        # Construct the normalized search string for Prd_name
        search_prd_name = clean_string(prd_name)

        found_image = None

        # Attempt to find matching image based on Prd_name
        for img_name, img_path in available_images.items():
            img_base_name = os.path.splitext(img_name)[0]
            if search_prd_name in clean_string(img_base_name):
                found_image = img_path
                break

        if found_image:
            try:
                # Construct the new filename in the format: prd_code Prd_name number
                new_file_name = f"{prd_code} {prd_name} {number}"

                # Add the original file extension to the new name
                file_extension = os.path.splitext(found_image)[1]
                new_file_path = os.path.join(os.path.dirname(found_image), f"{new_file_name}{file_extension}")

                # Rename the file
                os.rename(found_image, new_file_path)
                renamed_images.append((os.path.basename(found_image), new_file_name, new_file_path))
            except Exception as e:
                print(f"Error renaming image {os.path.basename(found_image)}: {e}")
        else:
            missing_images.append(prd_name)

    # Save missing images report if there are any missing images
    if missing_images:
        report_path = save_missing_images_to_word(missing_images, output_directory)
        messagebox.showinfo("Report Generated", f"Missing images report saved at:\n{report_path}")

    # Save renamed images log if any images were renamed
    if renamed_images:
        log_path = save_renamed_images_log(renamed_images, output_directory)
        messagebox.showinfo("Log Generated", f"Renamed images log saved at:\n{log_path}")

    # Display result messages
    if renamed_images:
        messagebox.showinfo("Success", f"Successfully renamed {len(renamed_images)} images!")
    else:
        messagebox.showwarning("No Images Renamed", "No images were renamed.")

# Create GUI Window
window = tk.Tk()
window.title("Image Renamer with Logs")

# Set window size
window.geometry("400x300")

# Add button to rename images
btn_rename = tk.Button(window, text="Rename Images", command=rename_images, width=20, height=2)
btn_rename.pack(pady=20)

# Run the GUI loop
window.mainloop()
